"""
文档加载器

支持多种文档格式:
- TXT
- PDF (使用 PaddleOCR API 解析扫描 PDF，返回 Markdown 格式)
- DOCX (优先使用 PaddleOCR，失败则使用 python-docx)
- Markdown
"""

from typing import List, Dict, Optional
from pathlib import Path
import os
import requests
import base64
import time
import tempfile
import shutil


# PaddleOCR API 配置
PADDLEOCR_API_URL = "https://feubu8f2jcses3d8.aistudio-app.com/layout-parsing"
PADDLEOCR_TOKEN = os.getenv("PADDLEOCR_TOKEN", "your_paddleocr_token_here")

API_TIMEOUT = 600
MAX_RETRIES = 3
RETRY_DELAY = 5


def call_paddleocr_api(file_path: str, max_retries: int = MAX_RETRIES) -> str:
    """调用 PaddleOCR API 解析 PDF"""
    with open(file_path, 'rb') as f:
        file_data = f.read()
    file_base64 = base64.b64encode(file_data).decode('utf-8')
    
    headers = {
        "Authorization": f"Bearer {PADDLEOCR_TOKEN}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "file": file_base64,
        "fileName": os.path.basename(file_path)
    }
    
    last_error = None
    
    for attempt in range(max_retries):
        try:
            response = requests.post(
                PADDLEOCR_API_URL,
                headers=headers,
                json=payload,
                timeout=API_TIMEOUT
            )
            
            if response.status_code == 503:
                last_error = "服务器繁忙"
                if attempt < max_retries - 1:
                    time.sleep(RETRY_DELAY)
                    continue
                raise Exception(last_error)
            
            if response.status_code == 400:
                raise Exception("文件格式不支持")
            
            if response.status_code != 200:
                raise Exception(f"API错误: {response.status_code}")
            
            result = response.json()
            
            texts = []
            if "result" in result and "layoutParsingResults" in result["result"]:
                for layout in result["result"]["layoutParsingResults"]:
                    if "markdown" in layout and isinstance(layout["markdown"], dict):
                        if "text" in layout["markdown"]:
                            text = layout["markdown"]["text"]
                            if text and isinstance(text, str):
                                texts.append(text)
            
            if texts:
                return "\n\n---\n\n".join(texts)
            return str(result.get("result", ""))
            
        except requests.exceptions.Timeout:
            last_error = "请求超时"
            if attempt < max_retries - 1:
                time.sleep(RETRY_DELAY)
        except requests.exceptions.RequestException as e:
            last_error = str(e)
            if attempt < max_retries - 1:
                time.sleep(RETRY_DELAY)
    
    raise Exception(last_error or "未知错误")


def extract_text_from_docx(file_path: str) -> str:
    """使用 python-docx 提取 DOCX 文本"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        return "\n\n".join(paragraphs)
        
    except ImportError:
        return ""
    except Exception as e:
        print(f"    提取失败: {e}")
        return ""


def split_markdown_with_langchain(text: str, source_name: str) -> List[Dict]:
    """使用 LangChain Markdown 分割器"""
    try:
        from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
        
        headers_to_split_on = [
            ("#", "header_1"),
            ("##", "header_2"),
            ("###", "header_3"),
            ("####", "header_4"),
        ]
        
        markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=headers_to_split_on,
            strip_headers=False
        )
        
        md_chunks = markdown_splitter.split_text(text)
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]
        )
        
        final_chunks = []
        for chunk in md_chunks:
            if len(chunk.page_content) > 1500:
                for sub in text_splitter.split_text(chunk.page_content):
                    final_chunks.append({
                        "content": sub,
                        "metadata": {**chunk.metadata, "source": source_name}
                    })
            else:
                final_chunks.append({
                    "content": chunk.page_content,
                    "metadata": {**chunk.metadata, "source": source_name}
                })
        
        return final_chunks
        
    except ImportError:
        return [{"content": p, "metadata": {"source": source_name}} 
                for p in text.split("\n\n") if p.strip()]


class DocumentLoader:
    """文档加载器"""
    
    @staticmethod
    def load_txt(file_path: str, encoding: str = "utf-8") -> Dict:
        """加载 TXT 文件"""
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        return {
            "content": content,
            "title": Path(file_path).stem,
            "metadata": {"source": file_path, "type": "txt"}
        }
    
    @staticmethod
    def load_markdown(file_path: str, encoding: str = "utf-8") -> Dict:
        """加载 Markdown 文件"""
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        
        title = Path(file_path).stem
        for line in content.split('\n'):
            if line.startswith('#'):
                title = line.lstrip('#').strip()
                break
        
        chunks = split_markdown_with_langchain(content, Path(file_path).name)
        return {
            "content": content,
            "title": title,
            "metadata": {"source": file_path, "type": "markdown"},
            "markdown_chunks": chunks
        }
    
    @staticmethod
    def load_pdf(file_path: str) -> Optional[Dict]:
        """加载 PDF 文件（使用 PaddleOCR API）"""
        file_name = Path(file_path).name
        file_size_kb = Path(file_path).stat().st_size / 1024
        
        print(f"  正在处理: {file_name} ({file_size_kb:.1f} KB)...")
        
        try:
            markdown_text = call_paddleocr_api(file_path)
            print(f"    成功! 提取 {len(markdown_text)} 字符")
            
            chunks = split_markdown_with_langchain(markdown_text, file_name)
            print(f"    分割为 {len(chunks)} 个块")
            
            return {
                "content": markdown_text,
                "title": Path(file_path).stem,
                "metadata": {
                    "source": file_path,
                    "type": "pdf",
                    "parser": "paddleocr_api"
                },
                "markdown_chunks": chunks
            }
        except Exception as e:
            print(f"    失败: {e}")
            return None
    
    @staticmethod
    def load_docx(file_path: str) -> Optional[Dict]:
        """
        加载 DOCX 文件
        
        策略：直接使用 python-docx 提取文本
        （因为 PaddleOCR 对 docx2pdf 生成的 PDF 支持不好）
        """
        file_name = Path(file_path).name
        file_size_kb = Path(file_path).stat().st_size / 1024
        
        print(f"  正在处理: {file_name} ({file_size_kb:.1f} KB)...")
        print(f"    使用 python-docx 提取文本...")
        
        content = extract_text_from_docx(file_path)
        
        if content:
            print(f"    成功! 提取 {len(content)} 字符")
            chunks = split_markdown_with_langchain(content, file_name)
            print(f"    分割为 {len(chunks)} 个块")
            
            return {
                "content": content,
                "title": Path(file_path).stem,
                "metadata": {
                    "source": file_path,
                    "type": "docx",
                    "parser": "python-docx"
                },
                "markdown_chunks": chunks
            }
        
        print(f"    失败: 无法提取文本")
        return None
    
    @classmethod
    def load_file(cls, file_path: str) -> Optional[Dict]:
        """根据文件类型加载文档"""
        ext = Path(file_path).suffix.lower()
        
        loaders = {
            '.txt': cls.load_txt,
            '.md': cls.load_markdown,
            '.pdf': cls.load_pdf,
            '.docx': cls.load_docx,
        }
        
        loader = loaders.get(ext)
        if loader:
            try:
                return loader(file_path)
            except Exception as e:
                print(f"加载失败 {file_path}: {e}")
                return None
        print(f"不支持的类型：{ext}")
        return None
    
    @classmethod
    def load_directory(
        cls,
        directory: str,
        extensions: List[str] = None
    ) -> List[Dict]:
        """加载目录下所有文档"""
        if extensions is None:
            extensions = ['.txt', '.md', '.pdf', '.docx']
        
        documents = []
        binary_files = []  # PDF 和 DOCX
        
        # 收集文件
        for root, _, files in os.walk(directory):
            for file in files:
                if Path(file).suffix.lower() in extensions:
                    file_path = os.path.join(root, file)
                    ext = Path(file).suffix.lower()
                    
                    if ext in ['.pdf', '.docx']:
                        binary_files.append(file_path)
                    else:
                        doc = cls.load_file(file_path)
                        if doc:
                            documents.append(doc)
        
        # 处理 PDF 和 DOCX
        if binary_files:
            print(f"\n正在处理 {len(binary_files)} 个文件 (PDF/DOCX)...")
            success_count = 0
            fail_count = 0
            
            for i, file_path in enumerate(binary_files, 1):
                print(f"\n[{i}/{len(binary_files)}]", end="")
                doc = cls.load_file(file_path)
                if doc:
                    documents.append(doc)
                    success_count += 1
                else:
                    fail_count += 1
            
            print(f"\n\n处理完成: 成功 {success_count}, 失败 {fail_count}")
        
        return documents
    
    @classmethod
    def get_markdown_chunks(cls, documents: List[Dict]) -> List[Dict]:
        """提取所有 Markdown 分块"""
        all_chunks = []
        for doc in documents:
            if doc and "markdown_chunks" in doc and doc["markdown_chunks"]:
                all_chunks.extend(doc["markdown_chunks"])
        return all_chunks